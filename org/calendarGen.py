# -*- coding: utf-8 -*-
"""
Created on Sun Mar 30 22:21:43 2025

@author: javit
"""

import calendar
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt  # Importar para manejar tamaños de fuente


meses_es = {
    "January": "Enero", "February": "Febrero", "March": "Marzo", "April": "Abril",
    "May": "Mayo", "June": "Junio", "July": "Julio", "August": "Agosto",
    "September": "Septiembre", "October": "Octubre", "November": "Noviembre", "December": "Diciembre"
}
font = 9
# Crear un documento con tabla de calendario mensual
doc_calendar = Document()
doc_calendar.add_heading('Calendario de Actividades - Abril y Mayo', 0)
section = doc_calendar.sections[0]
new_width = section.page_height  # Invertimos los valores de ancho y alto
new_height = section.page_width

section.page_width = new_width
section.page_height = new_height

# Crear la tabla mensual para abril
def add_month_table(doc, month, year, activities):
    # Agregar el nombre del mes en español
    nombre_mes = meses_es[calendar.month_name[month]]
    doc.add_heading(f'{nombre_mes} {year}', level=1)

    # Crear la tabla con 7 columnas (lunes, martes, miércoles, jueves, viernes, sábado, domingo)
    month_table = doc.add_table(rows=1, cols=7)
    days_of_week = ['Lun', 'Mar', 'Mié', 'Jue', 'Vie', 'Sáb', 'Dom']
    for i, day in enumerate(days_of_week):
        month_table.cell(0, i).text = day

    # Obtener el calendario del mes
    month_days = calendar.monthcalendar(year, month)

    for week in month_days:
        row = month_table.add_row().cells
        for i, day in enumerate(week):
            if day != 0:
                # Buscar actividad en la lista (con nombre del mes en español)
                day_str = str(day)
                activity = next((act for act in activities if act[0] == f"{day} de {meses_es[calendar.month_name[month]]}"), None)
                if activity:
                    p = row[i].paragraphs[0]  # Obtener el párrafo dentro de la celda
                    run = p.add_run(f"{day_str}\n{activity[1]}")
                    run.font.size = Pt(font)  # Cambiar el tamaño de fuente a 8 puntos

                else:
                    p = row[i].paragraphs[0]
                    run = p.add_run(day_str)
                    run.font.size = Pt(font)
            else:
                row[i].text = ''

    doc.add_paragraph("\n")  # Salto de línea para separar los meses
    
# Actividades para abril y mayo
actividades_abril_mayo = [
    ("31 de Marzo", "csvs de prueba a Albert"),
    ("31 de Marzo", "Preparar parser de datos usable modelo"),
    ("1 de Abril", "Preparar parser de datos usable modelo"),
    ("2 de Abril", "Preparar parser de datos usable modelo"),
    ("3 de Abril", "Preparar parser de datos usable modelo"),
    ("4 de Abril", "Preparar parser de datos usable modelo"),    
    ("11 de Abril", "Medidas Castelldefels"),
    ("14 de Abril", "Medidas CMP (IOMS+bolsasx10 a pie)"),
    ("15 de Abril", "Olfactometrias Odournet\n->Equipo IR GSS"),
    ("16 de Abril", "Semana Santa"),
    ("17 de Abril", "Semana Santa"),
    ("18 de Abril", "Semana Santa"),
    ("19 de Abril", "Semana Santa"),
    ("20 de Abril", "Semana Santa"),
    ("21 de Abril", "Semana Santa"),
    ("22 de Abril", "Medidas CMP (IOMS+bolsasx10 a pie)\n-> Olfactometrias Odournet\n-> Cover clases de Javi"),
    ("23 de Abril", "Test IR GSS"),
    ("24 de Abril", "Test IR GSS"),
    ("25 de Abril", "Test IR GSS"),
    ("7 de Mayo", "Medidas TRDB (IOMS+bolsasx10 vuelo)"),
    ("8 de Mayo", "Medidas TRDB (IOMS+bolsasx10 vuelo)\n-> Olfactomerias Odournet\n-> Cover clases de Javi"),
    ("9 de Mayo", "Olfactometrias Odournet"),
    ("21 de Mayo", "Medidas CMP (IOMS+bolsasx10 vuelo)"),
    ("22 de Mayo", "Medidas CMP (IOMS+bolsasx10 vuelo)\n-> Olfactometrias Odournet\n-> Cover clases de Javi"),
    ("23 de Mayo", "Olfactometrias Odournet"),
    ("4 de Junio", "Medidas TRDB (IOMS+bolsasx10 vuelo)"),
    ("5 de Junio", "Medidas TRDB (IOMS+bolsasx10 vuelo)\n-> Olfactomerias Odournet"),
    ("6 de Junio", "Olfactometrias Odournet"),
    ("18 de Junio", "Medidas FSL (IOMS+bolsasx10 a pie)"),
    ("19 de Junio", "Medidas FSL (IOMS+bolsasx10 a pie)\n-> Olfactomerias Odournet"),
    ("20 de Junio", "Olfactometrias Odournet"),

]

# Agregar tablas de calendario para abril y mayo
add_month_table(doc_calendar, 4, 2025, actividades_abril_mayo)  # Abril
add_month_table(doc_calendar, 5, 2025, actividades_abril_mayo)  # Mayo
add_month_table(doc_calendar, 6, 2025, actividades_abril_mayo)  # Mayo

# Guardar el documento con el calendario
file_path_calendar = "calendario_mensual_abril_mayo_Junio.docx"
doc_calendar.save(file_path_calendar)
